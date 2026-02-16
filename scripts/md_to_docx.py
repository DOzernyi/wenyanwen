#!/usr/bin/env python3
"""
Convert Markdown files with annotations to DOCX format.
Features:
  - 6"x9" paper, vertical Chinese text (tbRl)
  - 0.6" margins on all sides
  - 台灣明體 cwTeXMing font throughout
  - End-of-chapter footnotes grouped by label, with colors preserved
  - Proper Word heading styles (Heading 1/2/3)
  - Lists, indentation, beamer blocks with colors
"""

import os
import re
import sys
import glob
import copy
import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls

FONT_NAME = '台灣明體'
FONT_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)),
                         'assets', 'fonts', 'cwTeXMing', '台灣明體cwTeXMing.ttf')

HEADING1_SIZE = 17
HEADING2_SIZE = 15
HEADING3_SIZE = 12
NORMAL_SIZE = 13
FOOTNOTE_SIZE = 11

BLOCK_COLORS = {
    '': {'title_bg': '2a7ab5', 'body_bg': 'd9edf7', 'title_fg': 'FFFFFF'},
    'alert': {'title_bg': 'c0392b', 'body_bg': 'f2dede', 'title_fg': 'FFFFFF'},
    'example': {'title_bg': '27ae60', 'body_bg': 'dff0d8', 'title_fg': 'FFFFFF'},
    'warning': {'title_bg': 'e67e22', 'body_bg': 'fcf8e3', 'title_fg': 'FFFFFF'},
    'purple': {'title_bg': '8e44ad', 'body_bg': 'f0e6f6', 'title_fg': 'FFFFFF'},
    'note': {'title_bg': '0437F2', 'body_bg': 'e0e8ff', 'title_fg': 'FFFFFF'},
}


def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    if len(hex_color) == 3:
        hex_color = ''.join([c * 2 for c in hex_color])
    return RGBColor(int(hex_color[0:2], 16),
                    int(hex_color[2:4], 16),
                    int(hex_color[4:6], 16))


def set_run_font(run, size=NORMAL_SIZE, bold=False, color=None, font_name=FONT_NAME):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    if color:
        run.font.color.rgb = color


def set_paragraph_font(para, size=NORMAL_SIZE, font_name=FONT_NAME):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        para._element.insert(0, pPr)
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        pPr.append(rPr)
    sz = rPr.find(qn('w:sz'))
    if sz is not None:
        rPr.remove(sz)
    szCs = rPr.find(qn('w:szCs'))
    if szCs is not None:
        rPr.remove(szCs)
    rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{size * 2}"/>'))
    rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="{size * 2}"/>'))
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_shading(para_or_cell, color_hex):
    el = para_or_cell._element if hasattr(para_or_cell, '_element') else para_or_cell
    pPr = el.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        el.insert(0, pPr)
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    existing = pPr.find(qn('w:shd'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(shd)


def add_cell_shading(cell, color_hex):
    tcPr = cell._element.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
        cell._element.insert(0, tcPr)
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


# ─── Parsing ───────────────────────────────────────────────────────────────

FN_SENTINEL = '\x00FN:'
FN_END = '\x00:FN'
BLOCK_SENTINEL = '\x00BLOCK:'
BLOCK_END = '\x00:BLOCK'


def find_matching_endcapture(text, start_pos):
    depth = 1
    pos = start_pos
    cap_re = re.compile(r'\{%[-\s]*(?:capture\s+\w+|endcapture)\s*[-]?%\}')
    while depth > 0:
        m = cap_re.search(text, pos)
        if not m:
            return len(text)
        if 'endcapture' in m.group():
            depth -= 1
            if depth == 0:
                return m.start()
        else:
            depth += 1
        pos = m.end()
    return len(text)


def preprocess_liquid(md_content):
    footnotes = []
    blocks = []

    main_start_re = re.search(
        r'\{%[-\s]*capture\s+main_text\s*[-]?%\}', md_content
    )
    if main_start_re:
        body_start = main_start_re.end()
        body_end = find_matching_endcapture(md_content, body_start)
        content = md_content[body_start:body_end]
    else:
        content = md_content

    captures = {}
    inner_cap = re.compile(
        r'\{%[-\s]*capture\s+(\w+)\s*[-]?%\}(.*?)\{%[-\s]*endcapture\s*[-]?%\}',
        re.DOTALL
    )
    for m in inner_cap.finditer(content):
        name = m.group(1)
        if name != 'main_text':
            captures[name] = m.group(2).strip()

    content = inner_cap.sub('', content)

    def replace_fn(m):
        attrs = m.group(1)
        label = re.search(r'label="([^"]*)"', attrs)
        color = re.search(r'color="([^"]*)"', attrs)
        text = re.search(r'text="((?:[^"\\]|\\.)*)"', attrs)
        label_val = label.group(1) if label else '注'
        color_val = color.group(1) if color else '#267CB9'
        text_val = text.group(1) if text else ''
        if not text_val.strip():
            return ''
        idx = len(footnotes)
        footnotes.append({
            'label': label_val,
            'color': color_val,
            'text': text_val
        })
        return f'{FN_SENTINEL}{idx}{FN_END}'

    content = re.sub(
        r'\{%\s*include\s+fn\.html\s+(.*?)\s*%\}',
        replace_fn,
        content
    )

    def replace_block_include(m):
        attrs = m.group(1)
        type_m = re.search(r'type="([^"]*)"', attrs)
        title_m = re.search(r'title="([^"]*)"', attrs)
        content_m = re.search(r'content=(\w+)', attrs)
        block_type = type_m.group(1) if type_m else ''
        block_title = title_m.group(1) if title_m else ''
        block_var = content_m.group(1) if content_m else ''
        body = captures.get(block_var, '')
        body = re.sub(
            r'\{%\s*include\s+fn\.html\s+(.*?)\s*%\}',
            replace_fn,
            body
        )
        idx = len(blocks)
        blocks.append({
            'type': block_type,
            'title': block_title,
            'body': body
        })
        return f'\n\n{BLOCK_SENTINEL}{idx}{BLOCK_END}\n\n'

    content = re.sub(
        r'\{%\s*include\s+block\.html\s+(.*?)\s*%\}',
        replace_block_include,
        content
    )

    content = re.sub(r'\{%.*?%\}', '', content)
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)

    return content, footnotes, blocks


def convert_to_html(md_text):
    lines = md_text.split('\n')
    fixed = []
    for i, line in enumerate(lines):
        if re.match(r'\s*[\*\-]\s', line) and i > 0:
            prev = fixed[-1] if fixed else ''
            if prev.strip() and not re.match(r'\s*[\*\-]\s', prev) and not prev.strip().startswith('#'):
                fixed.append('')
        if re.match(r'\s*\d+\.\s', line) and i > 0:
            prev = fixed[-1] if fixed else ''
            if prev.strip() and not re.match(r'\s*\d+\.\s', prev) and not prev.strip().startswith('#'):
                fixed.append('')
        fixed.append(line)
    md_text = '\n'.join(fixed)
    html = markdown.markdown(md_text, extensions=['extra'])
    html = re.sub(r'<br\s*/?>', '<br/>', html)
    return html


def parse_content(md_content):
    content, footnotes, blocks = preprocess_liquid(md_content)

    title = ""
    front_matter = re.match(r'^---\s*\n(.*?\n)---\s*\n', md_content, re.DOTALL)
    if front_matter:
        title_m = re.search(r'title:\s*(.+)', front_matter.group(1))
        if title_m:
            title = title_m.group(1).strip()

    source_header = ''
    sh_match = re.search(r'<div\s+class="source-header">(.*?)</div>', md_content, re.DOTALL)
    if sh_match:
        source_header = re.sub(r'<[^>]+>', '', sh_match.group(1)).strip()

    parts = re.split(r'\x00BLOCK:(\d+)\x00:BLOCK', content)

    doc_elements = []
    if source_header:
        doc_elements.append({'type': 'source_header', 'text': source_header})

    for i, part in enumerate(parts):
        if i % 2 == 0:
            html = convert_to_html(part)
            elements = parse_html_to_elements(html, footnotes)
            doc_elements.extend(elements)
        else:
            block_idx = int(part)
            block = blocks[block_idx]
            body_html = convert_to_html(block['body'])
            body_elements = parse_html_to_elements(body_html, footnotes)
            doc_elements.append({
                'type': 'block',
                'block_type': block['type'],
                'title': block['title'],
                'body': body_elements
            })

    return title, doc_elements, footnotes


def parse_html_to_elements(html, footnotes):
    soup = BeautifulSoup(html, 'html.parser')
    elements = []

    for child in soup.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                elements.append({'type': 'paragraph', 'runs': parse_inline(child, footnotes)})
        elif isinstance(child, Tag):
            elements.extend(tag_to_elements(child, footnotes))

    return elements


def tag_to_elements(tag, footnotes, list_level=0):
    elements = []
    name = tag.name.lower() if tag.name else ''

    if name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        level = int(name[1])
        runs = parse_inline(tag, footnotes)
        elements.append({'type': 'heading', 'level': level, 'runs': runs})

    elif name == 'p':
        runs = parse_inline(tag, footnotes)
        if runs:
            elements.append({'type': 'paragraph', 'runs': runs})

    elif name == 'center':
        runs = parse_inline(tag, footnotes)
        if runs:
            elements.append({'type': 'paragraph', 'runs': runs, 'align': 'center'})

    elif name == 'blockquote':
        for child in tag.children:
            if isinstance(child, Tag):
                sub = tag_to_elements(child, footnotes, list_level)
                for el in sub:
                    el['indent'] = el.get('indent', 0) + 1
                elements.extend(sub)

    elif name in ('ul', 'ol'):
        ordered = (name == 'ol')
        item_num = 0
        for child in tag.children:
            if isinstance(child, Tag) and child.name == 'li':
                item_num += 1
                li_runs = []
                sub_lists = []
                for li_child in child.children:
                    if isinstance(li_child, Tag) and li_child.name in ('ul', 'ol'):
                        sub_lists.append(li_child)
                    else:
                        li_runs.extend(parse_inline(li_child, footnotes))
                if li_runs:
                    elements.append({
                        'type': 'list_item',
                        'ordered': ordered,
                        'level': list_level,
                        'number': item_num,
                        'runs': li_runs
                    })
                for sl in sub_lists:
                    elements.extend(tag_to_elements(sl, footnotes, list_level + 1))

    elif name == 'hr':
        elements.append({'type': 'hr'})

    elif name == 'pre':
        code = tag.get_text()
        elements.append({'type': 'code', 'text': code})

    elif name == 'div':
        for child in tag.children:
            if isinstance(child, Tag):
                elements.extend(tag_to_elements(child, footnotes, list_level))
            elif isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    elements.append({'type': 'paragraph', 'runs': [{'type': 'text', 'text': text}]})

    elif name == 'table':
        rows_data = []
        for section in [tag.find('thead'), tag.find('tbody'), tag]:
            if section:
                for tr in section.find_all('tr', recursive=False):
                    cells = []
                    for td in tr.find_all(['td', 'th'], recursive=False):
                        cells.append(parse_inline(td, footnotes))
                    if cells:
                        rows_data.append(cells)
            if rows_data:
                break
        if rows_data:
            elements.append({'type': 'table', 'rows': rows_data})

    else:
        runs = parse_inline(tag, footnotes)
        if runs:
            elements.append({'type': 'paragraph', 'runs': runs})

    return elements


def parse_inline(node, footnotes):
    runs = []
    if isinstance(node, NavigableString):
        text = str(node)
        _parse_text_with_sentinels(text, runs, footnotes)
    elif isinstance(node, Tag):
        name = node.name.lower() if node.name else ''
        if name == 'br':
            runs.append({'type': 'break'})
        elif name in ('strong', 'b'):
            for child in node.children:
                child_runs = parse_inline(child, footnotes)
                for r in child_runs:
                    r['bold'] = True
                runs.extend(child_runs)
        elif name in ('em', 'i'):
            for child in node.children:
                child_runs = parse_inline(child, footnotes)
                for r in child_runs:
                    r['italic'] = True
                runs.extend(child_runs)
        elif name in ('del', 's', 'strike'):
            for child in node.children:
                child_runs = parse_inline(child, footnotes)
                for r in child_runs:
                    r['strike'] = True
                runs.extend(child_runs)
        elif name == 'sup':
            for child in node.children:
                child_runs = parse_inline(child, footnotes)
                for r in child_runs:
                    r['superscript'] = True
                runs.extend(child_runs)
        elif name == 'sub':
            for child in node.children:
                child_runs = parse_inline(child, footnotes)
                for r in child_runs:
                    r['subscript'] = True
                runs.extend(child_runs)
        elif name == 'a':
            for child in node.children:
                runs.extend(parse_inline(child, footnotes))
        elif name == 'span':
            for child in node.children:
                runs.extend(parse_inline(child, footnotes))
        elif name == 'code':
            runs.append({'type': 'text', 'text': node.get_text(), 'code': True})
        else:
            for child in node.children:
                runs.extend(parse_inline(child, footnotes))
    return runs


def _parse_text_with_sentinels(text, runs, footnotes):
    pattern = re.compile(r'\x00FN:(\d+)\x00:FN')
    last = 0
    for m in pattern.finditer(text):
        before = text[last:m.start()]
        if before:
            runs.append({'type': 'text', 'text': before})
        idx = int(m.group(1))
        fn = footnotes[idx]
        runs.append({
            'type': 'footnote_ref',
            'index': idx,
            'label': fn['label'],
            'color': fn['color']
        })
        last = m.end()
    remainder = text[last:]
    if remainder:
        runs.append({'type': 'text', 'text': remainder})


# ─── DOCX Rendering ───────────────────────────────────────────────────────

def embed_font(doc, font_name, font_path):
    if not os.path.isfile(font_path):
        print(f"  Warning: Font file not found at {font_path}, skipping embed")
        return
    try:
        with open(font_path, 'rb') as f:
            font_data = f.read()

        part = doc.part
        font_part_name = f'/word/fonts/{font_name}.ttf'

        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        font_part = Part(
            PackURI(font_part_name),
            'application/x-font-ttf',
            font_data,
            part.package
        )

        part.relate_to(font_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/font')
    except Exception as e:
        print(f"  Note: Font embedding skipped ({e}), font must be installed locally")


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    sectPr = section._sectPr
    existing_td = sectPr.findall(qn('w:textDirection'))
    for td in existing_td:
        sectPr.remove(td)
    text_dir = parse_xml(
        f'<w:textDirection {nsdecls("w")} w:val="tbRl"/>'
    )
    sectPr.append(text_dir)

    embed_font(doc, FONT_NAME, FONT_PATH)
    setup_styles(doc)


def setup_styles(doc):
    styles = doc.styles

    for level, size in [(1, HEADING1_SIZE), (2, HEADING2_SIZE), (3, HEADING3_SIZE)]:
        style_name = f'Heading {level}'
        try:
            style = styles[style_name]
        except KeyError:
            from docx.enum.style import WD_STYLE_TYPE
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.name = FONT_NAME
        pf = style.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(4)
        el = style.element
        rPr = el.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            el.append(rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)

    try:
        normal_style = styles['Normal']
    except KeyError:
        pass
    else:
        normal_style.font.size = Pt(NORMAL_SIZE)
        normal_style.font.name = FONT_NAME
        nel = normal_style.element
        nrPr = nel.find(qn('w:rPr'))
        if nrPr is None:
            nrPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            nel.append(nrPr)
        nrFonts = nrPr.find(qn('w:rFonts'))
        if nrFonts is None:
            nrFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            nrPr.insert(0, nrFonts)
        nrFonts.set(qn('w:eastAsia'), FONT_NAME)
        nrFonts.set(qn('w:ascii'), FONT_NAME)
        nrFonts.set(qn('w:hAnsi'), FONT_NAME)


def render_elements(doc, elements, footnote_counters, size=NORMAL_SIZE):
    for el in elements:
        etype = el.get('type', '')

        if etype == 'source_header':
            para = doc.add_paragraph()
            run = para.add_run(el['text'])
            set_run_font(run, size=10, color=RGBColor(0xb0, 0xb0, 0xb0))
            para.paragraph_format.space_after = Pt(12)

        elif etype == 'heading':
            level = min(el['level'], 3)
            para = doc.add_paragraph(style=f'Heading {level}')
            render_runs(para, el['runs'], footnote_counters,
                        size=[0, HEADING1_SIZE, HEADING2_SIZE, HEADING3_SIZE][level],
                        bold=True)

        elif etype == 'paragraph':
            para = doc.add_paragraph()
            indent = el.get('indent', 0)
            if indent > 0:
                para.paragraph_format.left_indent = Inches(0.3 * indent)
            if el.get('align') == 'center':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            render_runs(para, el.get('runs', []), footnote_counters, size=size)

        elif etype == 'list_item':
            para = doc.add_paragraph()
            level = el.get('level', 0)
            para.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
            para.paragraph_format.first_line_indent = Inches(-0.2)
            ordered = el.get('ordered', False)
            num = el.get('number', 1)
            prefix = f'{num}. ' if ordered else '\u2022 '
            prefix_run = para.add_run(prefix)
            set_run_font(prefix_run, size=size)
            render_runs(para, el.get('runs', []), footnote_counters, size=size)

        elif etype == 'block':
            render_block(doc, el, footnote_counters)

        elif etype == 'hr':
            para = doc.add_paragraph()
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                para._element.insert(0, pPr)
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

        elif etype == 'code':
            para = doc.add_paragraph()
            run = para.add_run(el['text'])
            set_run_font(run, size=10)
            add_shading(para, 'F0F0F0')

        elif etype == 'table':
            rows_data = el.get('rows', [])
            if rows_data:
                num_cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=num_cols)
                table.style = 'Table Grid'
                for ri, row in enumerate(rows_data):
                    for ci, cell_runs in enumerate(row):
                        if ci < num_cols:
                            cell = table.rows[ri].cells[ci]
                            para = cell.paragraphs[0]
                            render_runs(para, cell_runs, footnote_counters, size=size)


def render_runs(para, runs, footnote_counters, size=NORMAL_SIZE, bold=False):
    for r in runs:
        if r.get('type') == 'text':
            text = r['text']
            if not text:
                continue
            run = para.add_run(text)
            set_run_font(run, size=size, bold=bold or r.get('bold', False))
            if r.get('italic'):
                run.font.italic = True
            if r.get('strike'):
                run.font.strike = True
            if r.get('superscript'):
                run.font.superscript = True
            if r.get('subscript'):
                run.font.subscript = True
            if r.get('code'):
                add_shading(para, 'F0F0F0')

        elif r.get('type') == 'footnote_ref':
            label = r['label']
            color_hex = r['color']
            if label not in footnote_counters:
                footnote_counters[label] = 0
            footnote_counters[label] += 1
            num = footnote_counters[label]
            marker_text = f'{label}{num}'
            run = para.add_run(marker_text)
            set_run_font(run, size=max(size - 2, 9), bold=True,
                         color=hex_to_rgb(color_hex))
            run.font.superscript = True

        elif r.get('type') == 'break':
            run = para.add_run()
            run.add_break()


def render_block(doc, block_el, footnote_counters):
    block_type = block_el.get('block_type', '')
    colors = BLOCK_COLORS.get(block_type, BLOCK_COLORS[''])
    title = block_el.get('title', '')

    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'

    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._element.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')

    title_cell = table.rows[0].cells[0]
    add_cell_shading(title_cell, colors['title_bg'])
    title_para = title_cell.paragraphs[0]
    title_run = title_para.add_run(title)
    set_run_font(title_run, size=NORMAL_SIZE, bold=True,
                 color=hex_to_rgb(colors['title_fg']))

    body_cell = table.rows[1].cells[0]
    add_cell_shading(body_cell, colors['body_bg'])

    body_elements = block_el.get('body', [])
    if body_elements:
        first = True
        for el in body_elements:
            if first:
                first_para = body_cell.paragraphs[0]
                if el.get('type') == 'paragraph':
                    render_runs(first_para, el.get('runs', []), footnote_counters,
                                size=NORMAL_SIZE)
                elif el.get('type') == 'list_item':
                    ordered = el.get('ordered', False)
                    num = el.get('number', 1)
                    prefix = f'{num}. ' if ordered else '\u2022 '
                    pr = first_para.add_run(prefix)
                    set_run_font(pr, size=NORMAL_SIZE)
                    render_runs(first_para, el.get('runs', []), footnote_counters,
                                size=NORMAL_SIZE)
                else:
                    render_runs(first_para, el.get('runs', []) if 'runs' in el else
                                [{'type': 'text', 'text': el.get('text', '')}],
                                footnote_counters, size=NORMAL_SIZE)
                first = False
                continue

            if el.get('type') == 'paragraph':
                p = body_cell.add_paragraph()
                render_runs(p, el.get('runs', []), footnote_counters, size=NORMAL_SIZE)
            elif el.get('type') == 'list_item':
                p = body_cell.add_paragraph()
                level = el.get('level', 0)
                p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                ordered = el.get('ordered', False)
                num = el.get('number', 1)
                prefix = f'{num}. ' if ordered else '\u2022 '
                pr = p.add_run(prefix)
                set_run_font(pr, size=NORMAL_SIZE)
                render_runs(p, el.get('runs', []), footnote_counters, size=NORMAL_SIZE)
            elif el.get('type') == 'heading':
                p = body_cell.add_paragraph()
                render_runs(p, el.get('runs', []), footnote_counters,
                            size=HEADING3_SIZE, bold=True)

    doc.add_paragraph()


def render_footnote_section(doc, footnotes, footnote_counters):
    if not footnotes:
        return

    grouped = {}
    label_order = []
    for fn in footnotes:
        label = fn['label']
        if label not in grouped:
            grouped[label] = []
            label_order.append(label)
        grouped[label].append(fn)

    active_groups = {l: fns for l, fns in grouped.items()
                     if l in footnote_counters and footnote_counters[l] > 0}
    if not active_groups:
        return

    hr_para = doc.add_paragraph()
    pPr = hr_para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        hr_para._element.insert(0, pPr)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    for label in label_order:
        if label not in active_groups:
            continue
        fns = active_groups[label]
        color = hex_to_rgb(fns[0]['color'])

        section_para = doc.add_paragraph()
        section_para.paragraph_format.space_before = Pt(8)
        section_run = section_para.add_run(f'── {label} ──')
        set_run_font(section_run, size=FOOTNOTE_SIZE + 1, bold=True, color=color)

        for i, fn in enumerate(fns, 1):
            fn_color = hex_to_rgb(fn['color'])
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)

            marker = para.add_run(f'{label}{i}')
            set_run_font(marker, size=FOOTNOTE_SIZE, bold=True, color=fn_color)

            sep = para.add_run('　')
            set_run_font(sep, size=FOOTNOTE_SIZE)

            text_run = para.add_run(fn['text'])
            set_run_font(text_run, size=FOOTNOTE_SIZE)


# ─── Main conversion ──────────────────────────────────────────────────────

def convert_md_to_docx(md_path, output_dir=None, source_dir=None):
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    title, elements, footnotes = parse_content(md_content)

    if not elements:
        print(f"  Skipping {md_path}: No content found")
        return None

    doc = Document()
    setup_document(doc)

    if title:
        title_para = doc.add_paragraph(style='Heading 1')
        title_run = title_para.add_run(title)
        set_run_font(title_run, size=HEADING1_SIZE, bold=True)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footnote_counters = {}
    render_elements(doc, elements, footnote_counters)

    render_footnote_section(doc, footnotes, footnote_counters)

    if output_dir is None:
        output_dir = os.path.dirname(md_path)
        rel_subdir = ''
    else:
        if source_dir:
            rel_path = os.path.relpath(os.path.dirname(md_path), source_dir)
            rel_subdir = rel_path if rel_path != '.' else ''
        else:
            rel_subdir = ''

    final_output_dir = os.path.join(output_dir, rel_subdir) if rel_subdir else output_dir
    os.makedirs(final_output_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(md_path))[0]
    output_path = os.path.join(final_output_dir, f"{base_name}.docx")
    doc.save(output_path)
    print(f"  -> {output_path}")
    return output_path


def convert_all_md_files(source_dir='pages', output_dir='output_docx'):
    os.makedirs(output_dir, exist_ok=True)
    md_files = sorted(glob.glob(os.path.join(source_dir, '**', '*.md'), recursive=True))

    converted = []
    for md_path in md_files:
        print(f"Converting: {md_path}")
        try:
            result = convert_md_to_docx(md_path, output_dir, source_dir)
            if result:
                converted.append(result)
        except Exception as e:
            print(f"  Error: {e}")
            import traceback
            traceback.print_exc()

    print(f"\nConverted {len(converted)} files to {output_dir}/")
    return converted


if __name__ == '__main__':
    if len(sys.argv) > 1:
        if sys.argv[1] == '--all':
            source = sys.argv[2] if len(sys.argv) > 2 else 'pages'
            output = sys.argv[3] if len(sys.argv) > 3 else 'output_docx'
            convert_all_md_files(source, output)
        else:
            for md_file in sys.argv[1:]:
                result = convert_md_to_docx(md_file)
                if result:
                    print(f"Created: {result}")
    else:
        print("Usage:")
        print("  python scripts/md_to_docx.py file.md")
        print("  python scripts/md_to_docx.py --all [source] [output]")
        print("")
        print("Examples:")
        print("  python scripts/md_to_docx.py pages/紅樓夢/紅樓夢1.md")
        print("  python scripts/md_to_docx.py --all pages output_docx")
